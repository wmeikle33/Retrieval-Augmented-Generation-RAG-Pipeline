from docx import Document
from dataclasses import dataclass
from typing import Dict


@dataclass
class ParsedDocument:
    doc_id: str
    text: str
    metadata: Dict


def parse_docx(path: str, doc_id: str) -> ParsedDocument:
    doc = Document(path)

    paragraphs = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if text:
            paragraphs.append(text)

    full_text = "\n".join(paragraphs)

    metadata = {
        "source": path,
        "paragraph_count": len(paragraphs)
    }

    return ParsedDocument(
        doc_id=doc_id,
        text=full_text,
        metadata=metadata
    )